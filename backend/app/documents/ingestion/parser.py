"""
DocuMind 2.0 — Document Parser
Uses unstructured.io to parse documents preserving structure.
"""

import os
from dataclasses import dataclass, field
from typing import Any

from loguru import logger


@dataclass
class ParsedElement:
    """A single parsed element from a document."""
    element_type: str  # Title, NarrativeText, Table, ListItem, Image, etc.
    text: str
    metadata: dict[str, Any] = field(default_factory=dict)
    is_table: bool = False
    table_html: str | None = None

    @property
    def page_number(self) -> int | None:
        return self.metadata.get("page_number")


@dataclass
class ParsedDocument:
    """A fully parsed document with typed elements."""
    file_path: str
    elements: list[ParsedElement] = field(default_factory=list)
    page_count: int = 0

    def add_element(
        self,
        element_type: str,
        text: str,
        metadata: dict[str, Any] | None = None,
    ) -> None:
        """Add a parsed element to the document."""
        meta = metadata or {}
        self.elements.append(
            ParsedElement(
                element_type=element_type,
                text=text,
                metadata=meta,
                is_table=meta.get("is_table", False),
                table_html=meta.get("table_html"),
            )
        )


async def parse_document(file_path: str, file_type: str | None = None) -> ParsedDocument:
    """
    Parse a document using unstructured.io, preserving structure.
    Handles PDF, DOCX, TXT, and other supported formats.

    Uses hi_res strategy for PDFs (OCR for scanned docs).
    Preserves tables, headings, lists, images, and narrative text.
    """
    if not os.path.exists(file_path):
        raise FileNotFoundError(f"Document not found: {file_path}")

    logger.info(f"Parsing document: {os.path.basename(file_path)} (type={file_type})")

    # Determine file type from extension if not provided
    if file_type is None:
        _, ext = os.path.splitext(file_path)
        file_type = ext.lower().lstrip(".")

    try:
        # Import unstructured here to allow graceful fallback
        from unstructured.partition.auto import partition
        from unstructured.documents.elements import Table

        # Configure parsing strategy based on file type
        partition_kwargs = {
            "filename": file_path,
            "include_page_breaks": True,
            "languages": ["eng"],
        }

        # Use hi_res for PDFs (enables OCR and table detection)
        if file_type == "pdf":
            partition_kwargs.update({
                "strategy": "hi_res",
                "infer_table_structure": True,
                "extract_images_in_pdf": False,  # Disable to avoid heavy deps initially
            })
        elif file_type in ("docx", "doc"):
            partition_kwargs.update({
                "strategy": "fast",
                "infer_table_structure": True,
            })
        else:
            partition_kwargs["strategy"] = "fast"

        # Parse the document
        elements = partition(**partition_kwargs)

        parsed = ParsedDocument(file_path=file_path)

        # Track max page number for page count
        max_page = 0

        for element in elements:
            page_num = getattr(element.metadata, "page_number", None)
            if page_num and page_num > max_page:
                max_page = page_num

            # Get table HTML if available
            table_html = None
            is_table = isinstance(element, Table)
            if is_table:
                table_html = getattr(element.metadata, "text_as_html", None)

            parsed.add_element(
                element_type=type(element).__name__,
                text=str(element),
                metadata={
                    "page_number": page_num,
                    "is_table": is_table,
                    "table_html": table_html,
                    "element_id": getattr(element, "id", None),
                    "source_file": os.path.basename(file_path),
                },
            )

        parsed.page_count = max_page or 1

        logger.info(
            f"Parsed {len(parsed.elements)} elements from {os.path.basename(file_path)} "
            f"({parsed.page_count} pages)"
        )

        return parsed

    except ImportError:
        logger.warning(
            "unstructured not available, falling back to basic text extraction"
        )
        return await _fallback_parse(file_path, file_type)


async def _fallback_parse(file_path: str, file_type: str) -> ParsedDocument:
    """
    Fallback parser when unstructured.io is not available.
    Uses pypdf for PDFs and basic text reading for other formats.
    """
    parsed = ParsedDocument(file_path=file_path)

    if file_type == "pdf":
        try:
            from pypdf import PdfReader

            reader = PdfReader(file_path)
            parsed.page_count = len(reader.pages)

            for i, page in enumerate(reader.pages, 1):
                text = page.extract_text()
                if text and text.strip():
                    parsed.add_element(
                        element_type="NarrativeText",
                        text=text.strip(),
                        metadata={
                            "page_number": i,
                            "is_table": False,
                            "source_file": os.path.basename(file_path),
                        },
                    )
        except Exception as e:
            logger.error(f"PDF parsing failed: {e}")
            raise

    elif file_type == "docx":
        try:
            from docx import Document as DocxDocument

            doc = DocxDocument(file_path)
            for i, para in enumerate(doc.paragraphs):
                if para.text.strip():
                    parsed.add_element(
                        element_type="NarrativeText",
                        text=para.text.strip(),
                        metadata={
                            "page_number": 1,
                            "is_table": False,
                            "source_file": os.path.basename(file_path),
                        },
                    )
            parsed.page_count = 1
        except Exception as e:
            logger.error(f"DOCX parsing failed: {e}")
            raise

    elif file_type == "txt":
        with open(file_path, "r", encoding="utf-8", errors="replace") as f:
            text = f.read()
        if text.strip():
            parsed.add_element(
                element_type="NarrativeText",
                text=text.strip(),
                metadata={
                    "page_number": 1,
                    "is_table": False,
                    "source_file": os.path.basename(file_path),
                },
            )
        parsed.page_count = 1

    else:
        raise ValueError(f"Unsupported file type: {file_type}")

    logger.info(
        f"Fallback parsed {len(parsed.elements)} elements from {os.path.basename(file_path)}"
    )
    return parsed
